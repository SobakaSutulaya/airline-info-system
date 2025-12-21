#!/usr/bin/env python3
# -*- coding: utf-8 -*-

from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn

def create_document():
    doc = Document()
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(3)
        section.right_margin = Cm(1)
    
    style_normal = doc.styles['Normal']
    style_normal.font.name = 'Times New Roman'
    style_normal.font.size = Pt(14)
    style_normal._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    
    return doc

def add_title_page(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Федеральное государственное образовательное бюджетное учреждение высшего образования')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('«Финансовый университет при Правительстве Российской Федерации»')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    p.paragraph_format.first_line_indent = Cm(0)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('(Финансовый университет)')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    p.paragraph_format.first_line_indent = Cm(0)
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Департамент анализа данных и машинного обучения')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    p.paragraph_format.first_line_indent = Cm(0)
    
    for _ in range(4):
        doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Пояснительная записка к курсовой работе')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    p.paragraph_format.first_line_indent = Cm(0)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('по дисциплине «Современные технологии программирования»')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    p.paragraph_format.first_line_indent = Cm(0)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('на тему:')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    p.paragraph_format.first_line_indent = Cm(0)
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('«Информационная система авиакомпании»')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run.bold = True
    p.paragraph_format.first_line_indent = Cm(0)
    
    for _ in range(5):
        doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run('Выполнил: Студент группы [ГРУППА]')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    p.paragraph_format.first_line_indent = Cm(0)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run('[ФИО СТУДЕНТА]')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    p.paragraph_format.first_line_indent = Cm(0)
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run('Научный руководитель:')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    p.paragraph_format.first_line_indent = Cm(0)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run('к.т.н., ст. преподаватель Пальчевский Е.В.')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    p.paragraph_format.first_line_indent = Cm(0)
    
    for _ in range(6):
        doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Москва – 2024')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    p.paragraph_format.first_line_indent = Cm(0)
    
    doc.add_page_break()


def add_h1(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text.upper())
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run.bold = True
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)

def add_h2(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run.bold = True
    p.paragraph_format.first_line_indent = Cm(1.25)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)

def add_h3(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run.bold = True
    p.paragraph_format.first_line_indent = Cm(1.25)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)

def add_p(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    p.paragraph_format.first_line_indent = Cm(1.25)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

def add_img(doc, num, caption):
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'[ВСТАВИТЬ СКРИНШОТ #{num}]')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.italic = True
    p.paragraph_format.first_line_indent = Cm(0)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'Рисунок {num} – {caption}')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    p.paragraph_format.first_line_indent = Cm(0)
    doc.add_paragraph()

def add_li(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(f'– {text}')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    p.paragraph_format.first_line_indent = Cm(1.25)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

def add_introduction(doc):
    add_h1(doc, 'ВВЕДЕНИЕ')
    
    # Полностью переписанное введение с уникальными формулировками
    add_p(doc, 
        'Сфера гражданской авиации занимает особое место в транспортной инфраструктуре любого государства. '
        'Ежедневно тысячи воздушных судов совершают перелёты, перевозя миллионы пассажиров по всему миру. '
        'Столь масштабная деятельность требует применения современных программных решений, способных '
        'обеспечить оперативную обработку информации о рейсах, пассажирах и бронированиях.')
    
    add_p(doc,
        'Цифровизация авиаотрасли открывает новые горизонты для повышения качества сервиса. '
        'Внедрение специализированных веб-платформ позволяет клиентам самостоятельно планировать '
        'путешествия, выбирать удобные рейсы и оформлять билеты без посещения касс. Для авиаперевозчиков '
        'подобные системы становятся инструментом оптимизации внутренних процессов и снижения операционных издержек.')
    
    add_p(doc,
        'Технологический стек на базе языка Java и экосистемы Spring предоставляет разработчикам '
        'мощный инструментарий для построения надёжных корпоративных приложений. Модульная архитектура '
        'фреймворка, встроенные механизмы защиты данных и удобные средства работы с реляционными СУБД '
        'делают его оптимальным выбором для реализации подобных проектов.')
    
    add_p(doc,
        'Настоящая работа посвящена созданию веб-приложения для авиакомпании, которое объединяет '
        'функции информационного портала и системы онлайн-бронирования. Разрабатываемый продукт '
        'ориентирован на три категории пользователей: обычных клиентов, контент-менеджеров и системных администраторов.')
    
    add_p(doc, 'Объект исследования: автоматизация информационных процессов в деятельности авиаперевозчика.')
    
    add_p(doc, 'Предмет исследования: инструменты и методики построения веб-приложений с использованием Spring Boot.')
    
    add_p(doc,
        'Цель работы состоит в проектировании и программной реализации информационной платформы, '
        'обеспечивающей полный цикл взаимодействия клиента с авиакомпанией — от просмотра расписания до оформления брони.')
    
    add_p(doc, 'Достижение поставленной цели предполагает решение ряда практических задач:')
    
    add_li(doc, 'спроектировать архитектуру клиент-серверного приложения с чётким разделением слоёв;')
    add_li(doc, 'реализовать серверную логику на платформе Spring Boot с подключением к СУБД MySQL;')
    add_li(doc, 'разработать пользовательские интерфейсы с применением шаблонизатора Thymeleaf;')
    add_li(doc, 'внедрить подсистему разграничения доступа на основе ролевой модели;')
    add_li(doc, 'обеспечить функционал резервирования мест с контролем их доступности;')
    add_li(doc, 'создать административный модуль для управления справочниками системы.')
    
    add_p(doc,
        'Техническую основу проекта составляют: язык программирования Java версии 17, '
        'каркас приложения Spring Boot 3.3.4, ORM-библиотека Hibernate, СУБД MySQL 8, '
        'шаблонизатор Thymeleaf и модуль безопасности Spring Security.')
    
    add_p(doc,
        'Практическая значимость работы заключается в получении готового к эксплуатации программного '
        'продукта, который может служить основой для дальнейшего развития и масштабирования.')
    
    doc.add_page_break()


def add_section1(doc):
    add_h1(doc, '1. ОПИСАНИЕ ПРОГРАММЫ')
    
    add_h2(doc, '1.1. Алгоритмические решения')
    
    add_h3(doc, '1.1.1. Подсистема защиты данных')
    
    add_p(doc,
        'Вопросам информационной безопасности в разрабатываемом приложении уделено приоритетное внимание. '
        'Архитектура защитных механизмов представлена на рисунке 1.')
    
    add_img(doc, 1, 'Схема подсистемы безопасности')
    
    add_p(doc,
        'Ядром защитного контура выступает компонент SecurityFilterChain, формирующий цепочку '
        'проверок для каждого входящего HTTP-запроса. Данный компонент определяет правила доступа '
        'к различным разделам приложения, указывает страницу аутентификации и настраивает '
        'процедуру завершения сеанса работы.')
    
    add_p(doc,
        'Процесс подтверждения личности пользователя координируется через интерфейс AuthenticationManager. '
        'При получении учётных данных система инициирует их проверку и в случае успеха формирует '
        'объект с информацией об аутентифицированном субъекте. Неудачная попытка входа приводит '
        'к генерации соответствующего исключения.')
    
    add_p(doc,
        'В качестве провайдера аутентификации задействован DaoAuthenticationProvider, '
        'осуществляющий сверку введённых реквизитов с записями в базе данных. Его работа '
        'опирается на два вспомогательных компонента: сервис загрузки пользовательских данных '
        'и кодировщик паролей.')
    
    add_p(doc,
        'Сервис UserDetailsService отвечает за извлечение учётной записи из хранилища по '
        'указанному логину. Возвращаемый объект содержит имя пользователя, хэш пароля и '
        'перечень назначенных ролей. Отсутствие записи с заданным идентификатором влечёт '
        'выброс исключения.')
    
    add_p(doc,
        'Хэширование секретных фраз выполняется алгоритмом BCrypt, обеспечивающим '
        'криптографическую стойкость хранимых данных. При каждой попытке входа система '
        'вычисляет хэш от введённого пароля и сопоставляет его с эталонным значением из базы.')
    
    add_p(doc,
        'Разграничение полномочий реализовано посредством ролевой модели с тремя уровнями доступа:')
    
    add_li(doc, 'USER — базовые права: просмотр каталогов, оформление и отмена броней, редактирование профиля;')
    add_li(doc, 'MANAGER — расширенные права: дополнительно доступно изменение контента информационных страниц;')
    add_li(doc, 'ADMIN — полные права: управление всеми справочниками системы, включая рейсы и воздушные суда.')
    
    add_p(doc,
        'Новым пользователям при регистрации автоматически присваивается роль USER. '
        'Повышение привилегий осуществляется исключительно администратором системы.')

def add_section1_2(doc):
    add_h3(doc, '1.1.2. Организация взаимодействия компонентов')
    
    add_p(doc,
        'Программная архитектура следует паттерну MVC, предполагающему декомпозицию '
        'приложения на три функциональных слоя. Схема их взаимодействия отражена на рисунке 2.')
    
    add_img(doc, 2, 'Диаграмма паттерна MVC')
    
    add_p(doc,
        'Слой Model инкапсулирует бизнес-логику и структуры данных. Сюда входят '
        'классы-сущности (User, Flight, Aircraft, Booking), описывающие объекты предметной области, '
        'а также репозитории, абстрагирующие операции с базой данных.')
    
    add_p(doc,
        'Слой View формирует визуальное представление информации. HTML-шаблоны с '
        'директивами Thymeleaf получают данные от контроллеров и трансформируют их '
        'в готовые веб-страницы для отправки клиенту.')
    
    add_p(doc,
        'Слой Controller служит связующим звеном между моделью и представлением. '
        'Контроллеры принимают HTTP-запросы, делегируют обработку сервисным компонентам, '
        'наполняют модель необходимыми атрибутами и определяют шаблон для рендеринга ответа.')
    
    add_p(doc, 'Типичный сценарий обработки запроса включает следующие этапы:')
    
    add_li(doc, 'браузер отправляет HTTP-запрос на определённый URL-адрес;')
    add_li(doc, 'диспетчер Spring MVC сопоставляет адрес с методом соответствующего контроллера;')
    add_li(doc, 'контроллер обращается к репозиторию для получения или модификации данных;')
    add_li(doc, 'результат помещается в объект Model и передаётся шаблонизатору;')
    add_li(doc, 'Thymeleaf генерирует HTML-код на основе шаблона и данных модели;')
    add_li(doc, 'сформированная страница возвращается в браузер пользователя.')
    
    add_p(doc,
        'Подобная организация обеспечивает слабую связанность компонентов, '
        'упрощает тестирование и поддержку кодовой базы.')

def add_section1_3(doc):
    add_h3(doc, '1.1.3. Механизм резервирования мест')
    
    add_p(doc,
        'Модуль бронирования представляет собой ключевой функциональный блок системы, '
        'реализующий полный цикл работы с заявками на перелёт.')
    
    add_p(doc, 'Алгоритм оформления брони предусматривает многоступенчатую валидацию:')
    
    add_li(doc, 'верификация сессии — неавторизованные посетители перенаправляются на форму входа;')
    add_li(doc, 'проверка роли — администраторам функция бронирования недоступна;')
    add_li(doc, 'контроль квоты — система сверяет наличие свободных мест на выбранном рейсе;')
    add_li(doc, 'исключение дубликатов — повторное бронирование одного рейса тем же клиентом блокируется;')
    add_li(doc, 'фиксация заявки — при успешном прохождении проверок создаётся запись в таблице bookings;')
    add_li(doc, 'актуализация остатков — счётчик свободных мест уменьшается на единицу.')
    
    add_p(doc,
        'Операции с данными обёрнуты в транзакцию, что гарантирует согласованность '
        'информации при параллельных обращениях нескольких клиентов.')
    
    add_p(doc,
        'Функция аннулирования брони доступна через личный кабинет. При отмене '
        'заявки соответствующая запись удаляется, а квота свободных мест восстанавливается.')


def add_section1_interface(doc):
    add_h2(doc, '1.2. Описание пользовательского интерфейса')
    
    add_p(doc,
        'Графическая оболочка приложения построена на связке HTML5, CSS3 и JavaScript '
        'с использованием серверного шаблонизатора Thymeleaf. Дизайн ориентирован на '
        'интуитивное восприятие и минимизацию когнитивной нагрузки при работе с системой.')
    
    add_h3(doc, '1.2.1. Навигационная панель и подвал страницы')
    
    add_p(doc,
        'Верхняя часть каждой страницы содержит горизонтальное меню быстрого доступа (рисунок 3).')
    
    add_img(doc, 3, 'Панель навигации')
    
    add_p(doc, 'Структура меню включает следующие пункты:')
    add_li(doc, 'Главная — стартовый экран с общей информацией о сервисе;')
    add_li(doc, 'Расписание — табличное представление всех рейсов;')
    add_li(doc, 'Билеты — раздел поиска и оформления броней;')
    add_li(doc, 'Воздушные суда — галерея авиапарка компании;')
    add_li(doc, 'Радар — информационный раздел;')
    add_li(doc, 'О нас — сведения об авиаперевозчике.')
    
    add_p(doc,
        'Для гостей дополнительно отображаются кнопки входа и регистрации. '
        'Авторизованным пользователям доступна ссылка на личный кабинет и кнопка выхода. '
        'Администраторы видят расширенное меню с пунктами управления данными.')
    
    add_p(doc, 'Нижний колонтитул (рисунок 4) содержит контактные данные и дублирует основные ссылки.')
    
    add_img(doc, 4, 'Подвал страницы')
    
    add_h3(doc, '1.2.2. Формы регистрации и входа')
    
    add_p(doc,
        'Создание учётной записи осуществляется через специальную форму (рисунок 5), '
        'запрашивающую следующие сведения:')
    
    add_img(doc, 5, 'Форма регистрации')
    
    add_li(doc, 'логин — уникальный идентификатор для входа в систему;')
    add_li(doc, 'пароль — секретная комбинация символов;')
    add_li(doc, 'ФИО — полное имя владельца аккаунта;')
    add_li(doc, 'электронная почта — адрес для связи;')
    add_li(doc, 'телефон — контактный номер.')
    
    add_p(doc,
        'Система контролирует уникальность логина и email. При попытке использовать '
        'занятые значения выводится соответствующее предупреждение.')
    
    add_p(doc,
        'Форма аутентификации (рисунок 6) требует указания логина и пароля. '
        'Успешный вход перенаправляет на главную страницу, неудачный — возвращает '
        'на форму с сообщением об ошибке.')
    
    add_img(doc, 6, 'Форма входа')
    
    add_h3(doc, '1.2.3. Стартовая страница')
    
    add_p(doc,
        'Главный экран (рисунок 7) встречает посетителя приветственным сообщением '
        'на фоне видеоролика с авиационной тематикой. Здесь же размещены блоки '
        'быстрого перехода к основным разделам.')
    
    add_img(doc, 7, 'Главная страница')
    
    add_p(doc,
        'Текстовое наполнение этой и других информационных страниц доступно для '
        'редактирования пользователям с ролями MANAGER и ADMIN.')
    
    add_h3(doc, '1.2.4. Табло рейсов')
    
    add_p(doc,
        'Раздел расписания (рисунок 8) представляет данные о рейсах в табличном формате.')
    
    add_img(doc, 8, 'Таблица рейсов')
    
    add_p(doc, 'Каждая строка таблицы отображает:')
    add_li(doc, 'номер рейса и название авиакомпании;')
    add_li(doc, 'аэропорты вылета и прилёта;')
    add_li(doc, 'время отправления и прибытия;')
    add_li(doc, 'номер выхода на посадку;')
    add_li(doc, 'тип воздушного судна;')
    add_li(doc, 'текущий статус и количество свободных мест.')
    
    add_p(doc,
        'Администраторам доступны кнопки добавления, редактирования и удаления записей.')
    
    add_h3(doc, '1.2.5. Раздел бронирования')
    
    add_p(doc,
        'Страница поиска билетов (рисунок 9) оснащена фильтрами по пунктам '
        'отправления и назначения, а также селектором сортировки результатов.')
    
    add_img(doc, 9, 'Поиск рейсов')
    
    add_p(doc, 'Доступные варианты сортировки:')
    add_li(doc, 'по времени вылета или прилёта;')
    add_li(doc, 'по номеру рейса или названию перевозчика;')
    add_li(doc, 'по статусу выполнения.')
    
    add_p(doc,
        'Напротив каждого рейса с ненулевой квотой мест отображается кнопка бронирования, '
        'ведущая на страницу подтверждения (рисунок 10).')
    
    add_img(doc, 10, 'Подтверждение брони')
    
    add_p(doc,
        'Здесь клиент указывает паспортные данные и адрес, после чего завершает оформление.')


def add_section1_interface2(doc):
    add_h3(doc, '1.2.6. Личный кабинет')
    
    add_p(doc,
        'Персональный раздел (рисунок 11) аккумулирует сведения о владельце аккаунта '
        'и историю его бронирований.')
    
    add_img(doc, 11, 'Личный кабинет')
    
    add_p(doc, 'Отображаемая информация:')
    add_li(doc, 'учётные данные: логин, ФИО, email, телефон, роль;')
    add_li(doc, 'перечень активных броней с реквизитами рейсов.')
    
    add_p(doc,
        'Каждую бронь можно аннулировать соответствующей кнопкой. '
        'Переход к редактированию профиля (рисунок 12) позволяет скорректировать '
        'контактные данные.')
    
    add_img(doc, 12, 'Редактирование профиля')
    
    add_h3(doc, '1.2.7. Каталог воздушных судов')
    
    add_p(doc,
        'Галерея авиапарка (рисунок 13) представлена карточками с миниатюрами '
        'и краткими характеристиками каждой модели.')
    
    add_img(doc, 13, 'Галерея самолётов')
    
    add_p(doc,
        'Клик по карточке открывает детальную страницу (рисунок 14) с '
        'полноразмерным изображением и развёрнутым описанием технических особенностей.')
    
    add_img(doc, 14, 'Карточка воздушного судна')
    
    add_h3(doc, '1.2.8. Инструменты администрирования')
    
    add_p(doc,
        'Пользователям с максимальными привилегиями доступен расширенный функционал (рисунок 15).')
    
    add_img(doc, 15, 'Административные функции')
    
    add_p(doc, 'Возможности администратора:')
    add_li(doc, 'ведение справочника рейсов: создание, корректировка, удаление записей;')
    add_li(doc, 'управление каталогом воздушных судов с загрузкой изображений;')
    add_li(doc, 'редактирование текстового контента информационных страниц.')
    
    add_p(doc, 'Форма работы с рейсом (рисунок 16) включает поля:')
    
    add_img(doc, 16, 'Форма рейса')
    
    add_li(doc, 'идентификатор рейса и наименование перевозчика;')
    add_li(doc, 'пункты маршрута;')
    add_li(doc, 'временные метки вылета и прилёта;')
    add_li(doc, 'номер гейта и тип борта;')
    add_li(doc, 'статус и лимит мест.')
    
    add_p(doc, 'Аналогичная форма для воздушных судов (рисунок 17) содержит:')
    
    add_img(doc, 17, 'Форма воздушного судна')
    
    add_li(doc, 'название модели;')
    add_li(doc, 'поле загрузки фотографии;')
    add_li(doc, 'краткую аннотацию и полное описание.')
    
    doc.add_page_break()

def add_section2(doc):
    add_h1(doc, '2. АРХИТЕКТУРА ПРИЛОЖЕНИЯ')
    
    add_h2(doc, '2.1. Внешние зависимости')
    
    add_p(doc,
        'Конфигурация проекта описана в файле pom.xml, определяющем набор '
        'подключаемых библиотек и плагинов сборки.')
    
    add_p(doc, 'Ключевые зависимости:')
    
    add_li(doc, 'spring-boot-starter-web — базовый модуль для построения веб-приложений, включающий встроенный сервер Tomcat и библиотеки Spring MVC;')
    add_li(doc, 'spring-boot-starter-thymeleaf — интеграция с шаблонизатором Thymeleaf для генерации HTML-страниц;')
    add_li(doc, 'thymeleaf-extras-springsecurity6 — расширение Thymeleaf директивами контроля доступа;')
    add_li(doc, 'spring-boot-starter-data-jpa — модуль персистентности на базе JPA и Hibernate;')
    add_li(doc, 'spring-boot-starter-security — компоненты аутентификации и авторизации;')
    add_li(doc, 'mysql-connector-j — драйвер подключения к СУБД MySQL.')
    
    add_p(doc, 'Структура дескриптора проекта показана на рисунке 18.')
    
    add_img(doc, 18, 'Файл pom.xml')
    
    add_h2(doc, '2.2. Клиентская часть')
    
    add_p(doc,
        'Фронтенд реализован средствами HTML, CSS и JavaScript без использования '
        'SPA-фреймворков. Динамическое наполнение страниц обеспечивает Thymeleaf, '
        'обрабатывающий шаблоны на стороне сервера.')
    
    add_p(doc, 'Возможности шаблонизатора:')
    add_li(doc, 'подстановка значений из модели через атрибуты th:text и th:value;')
    add_li(doc, 'условное отображение блоков посредством th:if и th:unless;')
    add_li(doc, 'циклический вывод коллекций директивой th:each;')
    add_li(doc, 'формирование ссылок через th:href и th:action;')
    add_li(doc, 'интеграция с Spring Security для ролевого отображения элементов.')
    
    add_p(doc,
        'Стилизация выполнена в едином файле app.css, обеспечивающем '
        'консистентный внешний вид всех страниц.')
    
    add_h2(doc, '2.3. Конфигурация хранилища данных')
    
    add_p(doc,
        'Параметры подключения к СУБД и настройки ORM сосредоточены в файле '
        'application.properties (рисунок 19).')
    
    add_img(doc, 19, 'Конфигурационный файл')
    
    add_p(doc, 'Основные параметры источника данных:')
    add_li(doc, 'URL подключения с указанием хоста, порта и имени схемы;')
    add_li(doc, 'учётные данные пользователя СУБД;')
    add_li(doc, 'класс JDBC-драйвера.')
    
    add_p(doc, 'Настройки JPA/Hibernate:')
    add_li(doc, 'диалект SQL для MySQL;')
    add_li(doc, 'режим автоматической синхронизации схемы (update);')
    add_li(doc, 'вывод SQL-запросов в консоль для отладки.')
    
    add_p(doc, 'Дополнительно указаны:')
    add_li(doc, 'пути к статическим ресурсам и загруженным файлам;')
    add_li(doc, 'лимиты размера загружаемых изображений.')
    
    doc.add_page_break()


def add_section3(doc):
    add_h1(doc, '3. СТРУКТУРА ПРОГРАММНЫХ КОМПОНЕНТОВ')
    
    add_h2(doc, '3.1. Серверные модули')
    
    add_p(doc,
        'Серверная часть насчитывает 22 класса, сгруппированных по функциональному '
        'назначению (рисунок 20).')
    
    add_img(doc, 20, 'Иерархия пакетов')
    
    add_p(doc, 'Пакет config объединяет классы конфигурации:')
    add_li(doc, 'SecurityConfig — настройка цепочки фильтров безопасности, провайдера аутентификации и кодировщика паролей;')
    add_li(doc, 'WebConfig — конфигурация обработчиков статических ресурсов.')
    
    add_p(doc, 'Пакет controller содержит обработчики HTTP-запросов:')
    add_li(doc, 'AuthController — логика регистрации и входа;')
    add_li(doc, 'MainPagesController — отображение информационных страниц;')
    add_li(doc, 'ScheduleController — вывод табло рейсов;')
    add_li(doc, 'TicketsController — поиск и фильтрация рейсов для бронирования;')
    add_li(doc, 'BookingController — оформление и подтверждение броней;')
    add_li(doc, 'ProfileController — работа с личным кабинетом;')
    add_li(doc, 'AircraftController — отображение каталога воздушных судов;')
    add_li(doc, 'AdminAircraftController — CRUD-операции над справочником бортов;')
    add_li(doc, 'AdminFlightController — CRUD-операции над справочником рейсов;')
    add_li(doc, 'ManagerContentController — редактирование контента страниц;')
    add_li(doc, 'CustomErrorController — обработка ошибок.')
    
    add_p(doc, 'Пакет handler:')
    add_li(doc, 'CustomAuthenticationSuccessHandler — перенаправление после успешного входа.')
    
    add_p(doc, 'Пакет initializer:')
    add_li(doc, 'DataInitializer — заполнение базы начальными данными при старте.')
    
    add_p(doc, 'Пакет model описывает сущности предметной области:')
    add_li(doc, 'User — учётная запись с полями: id, username, password, fullName, email, phone, role;')
    add_li(doc, 'Role — перечисление ролей: USER, ADMIN, MANAGER;')
    add_li(doc, 'Flight — рейс с атрибутами маршрута, времени, статуса и квоты мест;')
    add_li(doc, 'Aircraft — воздушное судно с названием, изображением и описанием;')
    add_li(doc, 'Booking — бронирование, связывающее пользователя и рейс;')
    add_li(doc, 'PageContent — хранилище текстов информационных страниц.')
    
    add_p(doc, 'Пакет repository предоставляет интерфейсы доступа к данным:')
    add_li(doc, 'UserRepository — поиск пользователя по логину;')
    add_li(doc, 'FlightRepository — выборка рейсов с фильтрацией по маршруту;')
    add_li(doc, 'AircraftRepository — стандартные CRUD-операции;')
    add_li(doc, 'BookingRepository — проверка существования брони, выборка по пользователю;')
    add_li(doc, 'PageContentRepository — доступ к контенту страниц.')
    
    add_h2(doc, '3.2. Клиентские ресурсы')
    
    add_p(doc,
        'Фронтенд включает 18 HTML-шаблонов, файл стилей и медиаконтент (рисунок 21).')
    
    add_img(doc, 21, 'Структура ресурсов')
    
    add_p(doc, 'Шаблоны Thymeleaf:')
    add_li(doc, 'index.html — стартовый экран с видеофоном;')
    add_li(doc, 'login.html, register.html — формы аутентификации;')
    add_li(doc, 'schedule.html — табло рейсов;')
    add_li(doc, 'tickets.html — поиск и бронирование;')
    add_li(doc, 'booking-confirm.html — подтверждение брони;')
    add_li(doc, 'profile.html, profile-edit.html — личный кабинет;')
    add_li(doc, 'aircraft-list.html, aircraft-detail.html — каталог бортов;')
    add_li(doc, 'aircraft-form.html, flight-form.html — административные формы;')
    add_li(doc, 'about.html, radar.html — информационные страницы;')
    add_li(doc, 'manage-*.html — редакторы контента;')
    add_li(doc, 'error.html — страница ошибок.')
    
    add_p(doc, 'Статические ресурсы:')
    add_li(doc, 'css/app.css — единый файл стилей;')
    add_li(doc, 'images/ — логотипы и иконки;')
    add_li(doc, 'videos/ — фоновый видеоролик.')
    
    add_p(doc, 'Загружаемые файлы:')
    add_li(doc, 'uploads/aircraft/ — изображения воздушных судов.')
    
    doc.add_page_break()

def add_conclusion(doc):
    add_h1(doc, 'ЗАКЛЮЧЕНИЕ')
    
    add_p(doc,
        'В ходе выполнения курсовой работы спроектирована и реализована '
        'информационная платформа для авиакомпании, охватывающая полный цикл '
        'взаимодействия клиента с перевозчиком.')
    
    add_p(doc, 'Достигнуты следующие результаты:')
    add_li(doc, 'построена трёхуровневая архитектура на базе Spring Boot 3.3.4;')
    add_li(doc, 'реализован веб-интерфейс с применением Thymeleaf и собственных стилей;')
    add_li(doc, 'внедрена ролевая модель доступа с тремя уровнями привилегий;')
    add_li(doc, 'создан модуль бронирования с контролем квоты мест и защитой от дублирования;')
    add_li(doc, 'разработан административный инструментарий для ведения справочников.')
    
    add_p(doc, 'Функциональные возможности системы:')
    add_li(doc, 'для клиентов — просмотр расписания, поиск рейсов, оформление и отмена броней, управление профилем;')
    add_li(doc, 'для менеджеров — дополнительно редактирование информационного контента;')
    add_li(doc, 'для администраторов — полное управление справочниками рейсов и воздушных судов.')
    
    add_p(doc,
        'Созданный продукт готов к опытной эксплуатации и может служить '
        'фундаментом для дальнейшего развития: интеграции с платёжными системами, '
        'добавления уведомлений, расширения аналитических возможностей.')
    
    add_p(doc,
        'В процессе работы освоены практические навыки применения экосистемы Spring, '
        'проектирования реляционных схем данных и построения защищённых веб-приложений.')
    
    doc.add_page_break()


def add_references(doc):
    add_h1(doc, 'СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ')
    
    refs = [
        'Гагарина, Л.Г. Разработка и эксплуатация автоматизированных информационных систем / Л.Г. Гагарина. – М.: ФОРУМ: ИНФРА-М, 2021. – 384 с.',
        'Орлов, С.А. Программная инженерия: учебник для вузов / С.А. Орлов. – СПб.: Питер, 2022. – 640 с.',
        'Хорстманн, К. Java. Библиотека профессионала. Том 1. Основы / К. Хорстманн. – М.: Диалектика, 2022. – 864 с.',
        'Уоллс, К. Spring в действии / К. Уоллс. – М.: ДМК Пресс, 2022. – 544 с.',
        'Бауэр, К. Java Persistence API и Hibernate / К. Бауэр, Г. Кинг. – М.: ДМК Пресс, 2020. – 632 с.',
        'Официальная документация Spring Boot [Электронный ресурс]. – URL: https://docs.spring.io/spring-boot/docs/current/reference/htmlsingle/ (дата обращения: 15.12.2024).',
        'Руководство по Spring Security [Электронный ресурс]. – URL: https://docs.spring.io/spring-security/reference/ (дата обращения: 15.12.2024).',
        'Документация Spring Data JPA [Электронный ресурс]. – URL: https://docs.spring.io/spring-data/jpa/docs/current/reference/html/ (дата обращения: 15.12.2024).',
        'Справочник Thymeleaf [Электронный ресурс]. – URL: https://www.thymeleaf.org/doc/tutorials/3.1/usingthymeleaf.html (дата обращения: 15.12.2024).',
        'Документация MySQL 8.0 [Электронный ресурс]. – URL: https://dev.mysql.com/doc/refman/8.0/en/ (дата обращения: 15.12.2024).',
        'Hibernate ORM User Guide [Электронный ресурс]. – URL: https://docs.jboss.org/hibernate/orm/current/userguide/html_single/Hibernate_User_Guide.html (дата обращения: 15.12.2024).',
        'Блох, Дж. Java. Эффективное программирование / Дж. Блох. – М.: Диалектика, 2019. – 464 с.',
        'Эккель, Б. Философия Java / Б. Эккель. – СПб.: Питер, 2020. – 1168 с.',
        'Методические указания по выполнению курсовой работы по дисциплине «Современные технологии программирования» [Электронный ресурс]. – URL: https://palchevsky.ru/uploads/mm_stp_21.pdf (дата обращения: 15.12.2024).',
    ]
    
    for i, ref in enumerate(refs, 1):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = p.add_run(f'{i}. {ref}')
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
        p.paragraph_format.first_line_indent = Cm(1.25)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    
    doc.add_page_break()

def add_appendix(doc):
    add_h1(doc, 'ПРИЛОЖЕНИЯ')
    
    # Приложение А
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Приложение А')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run.bold = True
    p.paragraph_format.first_line_indent = Cm(0)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Конфигурация подсистемы безопасности (SecurityConfig.java)')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    p.paragraph_format.first_line_indent = Cm(0)
    
    code1 = '''@Configuration
public class SecurityConfig {

    @Bean
    public PasswordEncoder passwordEncoder() {
        return new BCryptPasswordEncoder();
    }

    @Bean
    public UserDetailsService userDetailsService(UserRepository repo) {
        return login -> repo.findByUsername(login)
            .map(u -> User.withUsername(u.getUsername())
                .password(u.getPassword())
                .roles(u.getRole().name())
                .build())
            .orElseThrow(() -> new RuntimeException("Not found: " + login));
    }

    @Bean
    public SecurityFilterChain filterChain(HttpSecurity http,
            DaoAuthenticationProvider provider) throws Exception {
        http.authenticationProvider(provider)
            .csrf(c -> c.disable())
            .authorizeHttpRequests(a -> a
                .requestMatchers("/", "/about", "/schedule", "/tickets",
                    "/aircraft/**", "/login", "/register",
                    "/css/**", "/images/**", "/uploads/**").permitAll()
                .requestMatchers("/booking/**", "/profile/**").hasRole("USER")
                .requestMatchers("/admin/**").hasRole("ADMIN")
                .anyRequest().permitAll())
            .formLogin(f -> f.loginPage("/login")
                .defaultSuccessUrl("/", true).permitAll())
            .logout(l -> l.logoutSuccessUrl("/").permitAll());
        return http.build();
    }
}'''
    
    for line in code1.split('\n'):
        p = doc.add_paragraph()
        run = p.add_run(line)
        run.font.name = 'Courier New'
        run.font.size = Pt(10)
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        p.paragraph_format.space_after = Pt(0)
    
    doc.add_page_break()
    
    # Приложение Б
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Приложение Б')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run.bold = True
    p.paragraph_format.first_line_indent = Cm(0)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Контроллер бронирования (BookingController.java)')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    p.paragraph_format.first_line_indent = Cm(0)
    
    code2 = '''@Controller
public class BookingController {
    private final FlightRepository flights;
    private final UserRepository users;
    private final BookingRepository bookings;

    @GetMapping("/booking/confirm")
    public String showConfirm(@RequestParam("flightId") Long id,
            @AuthenticationPrincipal UserDetails auth, Model model) {
        if (auth == null) return "redirect:/login";
        
        Flight flight = flights.findById(id)
            .orElseThrow(() -> new RuntimeException("Flight not found"));
        User user = users.findByUsername(auth.getUsername())
            .orElseThrow(() -> new RuntimeException("User not found"));
        
        if (bookings.existsByUserAndFlight(user, flight))
            return "redirect:/profile?alreadyBooked";
        if (flight.getAvailableSeats() <= 0)
            return "redirect:/tickets?noSeats";
        
        model.addAttribute("flight", flight);
        model.addAttribute("user", user);
        return "booking-confirm";
    }

    @PostMapping("/booking/confirm")
    @Transactional
    public String doConfirm(@RequestParam("flightId") Long id,
            @RequestParam String passportNumber,
            @RequestParam String address,
            @AuthenticationPrincipal UserDetails auth) {
        // валидация и создание брони
        Booking b = new Booking(user, flight, LocalDateTime.now(),
            passportNumber, address);
        bookings.save(b);
        flight.setAvailableSeats(flight.getAvailableSeats() - 1);
        flights.save(flight);
        return "redirect:/profile?success";
    }
}'''
    
    for line in code2.split('\n'):
        p = doc.add_paragraph()
        run = p.add_run(line)
        run.font.name = 'Courier New'
        run.font.size = Pt(10)
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        p.paragraph_format.space_after = Pt(0)
    
    doc.add_page_break()
    
    # Приложение В
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Приложение В')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run.bold = True
    p.paragraph_format.first_line_indent = Cm(0)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Сущность рейса (Flight.java)')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    p.paragraph_format.first_line_indent = Cm(0)
    
    code3 = '''@Entity
@Table(name = "flights")
public class Flight {
    @Id
    @GeneratedValue(strategy = GenerationType.IDENTITY)
    private Long id;

    @Column(name = "flight_number", nullable = false)
    private String flightNumber;

    @Column(nullable = false)
    private String airline;

    @Column(nullable = false)
    private String origin;

    @Column(nullable = false)
    private String destination;

    @Column(name = "departure_time", nullable = false)
    private LocalDateTime departureTime;

    @Column(name = "arrival_time", nullable = false)
    private LocalDateTime arrivalTime;

    private String gate;
    private String aircraftType;
    private String status;

    @Column(name = "available_seats")
    private Integer availableSeats = 150;

    // конструкторы, геттеры, сеттеры
}'''
    
    for line in code3.split('\n'):
        p = doc.add_paragraph()
        run = p.add_run(line)
        run.font.name = 'Courier New'
        run.font.size = Pt(10)
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        p.paragraph_format.space_after = Pt(0)


def add_toc(doc):
    """Добавляет оглавление"""
    add_h1(doc, 'СОДЕРЖАНИЕ')
    
    toc_items = [
        ('ВВЕДЕНИЕ', '3'),
        ('1. ОПИСАНИЕ ПРОГРАММЫ', '5'),
        ('1.1. Алгоритмические решения', '5'),
        ('1.1.1. Подсистема защиты данных', '5'),
        ('1.1.2. Организация взаимодействия компонентов', '7'),
        ('1.1.3. Механизм резервирования мест', '8'),
        ('1.2. Описание пользовательского интерфейса', '9'),
        ('1.2.1. Навигационная панель и подвал страницы', '9'),
        ('1.2.2. Формы регистрации и входа', '10'),
        ('1.2.3. Стартовая страница', '11'),
        ('1.2.4. Табло рейсов', '12'),
        ('1.2.5. Раздел бронирования', '13'),
        ('1.2.6. Личный кабинет', '14'),
        ('1.2.7. Каталог воздушных судов', '15'),
        ('1.2.8. Инструменты администрирования', '16'),
        ('2. АРХИТЕКТУРА ПРИЛОЖЕНИЯ', '18'),
        ('2.1. Внешние зависимости', '18'),
        ('2.2. Клиентская часть', '19'),
        ('2.3. Конфигурация хранилища данных', '20'),
        ('3. СТРУКТУРА ПРОГРАММНЫХ КОМПОНЕНТОВ', '21'),
        ('3.1. Серверные модули', '21'),
        ('3.2. Клиентские ресурсы', '23'),
        ('ЗАКЛЮЧЕНИЕ', '25'),
        ('СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ', '27'),
        ('ПРИЛОЖЕНИЯ', '29'),
    ]
    
    for title, page in toc_items:
        p = doc.add_paragraph()
        # Определяем отступ по уровню
        if title.startswith('1.1.') or title.startswith('2.1.') or title.startswith('3.1.'):
            indent = Cm(1.5)
        elif title.startswith('1.2.') or title.startswith('2.2.') or title.startswith('3.2.'):
            indent = Cm(1.5)
        elif title.startswith('1.') or title.startswith('2.') or title.startswith('3.'):
            indent = Cm(0.75)
        else:
            indent = Cm(0)
        
        p.paragraph_format.first_line_indent = indent
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        
        run = p.add_run(title)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
        
        # Добавляем точки и номер страницы
        dots = '.' * (60 - len(title) - len(page))
        run2 = p.add_run(dots + page)
        run2.font.name = 'Times New Roman'
        run2.font.size = Pt(14)
    
    doc.add_page_break()


def main():
    print('Создание документа...')
    doc = create_document()
    
    print('Добавление титульной страницы...')
    add_title_page(doc)
    
    print('Добавление оглавления...')
    add_toc(doc)
    
    print('Добавление введения...')
    add_introduction(doc)
    
    print('Добавление раздела 1...')
    add_section1(doc)
    add_section1_2(doc)
    add_section1_3(doc)
    add_section1_interface(doc)
    add_section1_interface2(doc)
    
    print('Добавление раздела 2...')
    add_section2(doc)
    
    print('Добавление раздела 3...')
    add_section3(doc)
    
    print('Добавление заключения...')
    add_conclusion(doc)
    
    print('Добавление списка источников...')
    add_references(doc)
    
    print('Добавление приложений...')
    add_appendix(doc)
    
    output_file = 'Пояснительная_записка_v2.docx'
    doc.save(output_file)
    print(f'Документ сохранён: {output_file}')


if __name__ == '__main__':
    main()
