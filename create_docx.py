#!/usr/bin/env python3
# -*- coding: utf-8 -*-

from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_margins(cell, top=0, start=0, bottom=0, end=0):
    """Установка отступов в ячейке"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for margin_name, margin_value in [('top', top), ('start', start), ('bottom', bottom), ('end', end)]:
        node = OxmlElement(f'w:{margin_name}')
        node.set(qn('w:w'), str(margin_value))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def create_document():
    doc = Document()
    
    # Настройка полей страницы по ГОСТу
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(3)
        section.right_margin = Cm(1)
        section.header_distance = Cm(2)
        section.footer_distance = Cm(1.25)
    
    # Создание стилей
    styles = doc.styles
    
    # Стиль для обычного текста
    style_normal = styles['Normal']
    style_normal.font.name = 'Times New Roman'
    style_normal.font.size = Pt(14)
    style_normal._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    style_normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    style_normal.paragraph_format.first_line_indent = Cm(1.25)
    style_normal.paragraph_format.space_after = Pt(0)
    style_normal.paragraph_format.space_before = Pt(0)
    
    return doc

def add_title_page(doc):
    """Титульная страница"""
    # Шапка
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Федеральное государственное образовательное бюджетное учреждение высшего образования')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('«Финансовый университет')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('при Правительстве Российской Федерации» (Финансовый университет)')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Департамент анализа данных и машинного обучения')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    p.paragraph_format.first_line_indent = Cm(0)
    
    # Пустые строки
    for _ in range(4):
        doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Пояснительная записка к курсовой работе по дисциплине')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    p.paragraph_format.first_line_indent = Cm(0)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('«Современные технологии программирования» на тему:')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    p.paragraph_format.first_line_indent = Cm(0)
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('«Информационная система авиакомпании»')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run.bold = True
    p.paragraph_format.first_line_indent = Cm(0)
    
    # Пустые строки
    for _ in range(5):
        doc.add_paragraph()
    
    # Информация о студенте (справа)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run('Выполнил: Студент группы [ГРУППА]')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    p.paragraph_format.first_line_indent = Cm(0)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run('[ФИО СТУДЕНТА]')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    p.paragraph_format.first_line_indent = Cm(0)
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run('Научный руководитель: к.т.н., ст. преподаватель')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    p.paragraph_format.first_line_indent = Cm(0)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run('Пальчевский Е.В.')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    p.paragraph_format.first_line_indent = Cm(0)
    
    # Пустые строки до низа
    for _ in range(6):
        doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Москва – 2024')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    p.paragraph_format.first_line_indent = Cm(0)
    
    # Разрыв страницы
    doc.add_page_break()

def add_heading1(doc, text):
    """Заголовок первого уровня (ВВЕДЕНИЕ, ЗАКЛЮЧЕНИЕ и т.д.)"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text.upper())
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run.bold = True
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

def add_heading2(doc, text):
    """Заголовок второго уровня (1.1. Название)"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run.bold = True
    p.paragraph_format.first_line_indent = Cm(1.25)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

def add_heading3(doc, text):
    """Заголовок третьего уровня (1.1.1. Название)"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run.bold = True
    p.paragraph_format.first_line_indent = Cm(1.25)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

def add_paragraph_text(doc, text, first_indent=True):
    """Обычный абзац текста"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    p.paragraph_format.first_line_indent = Cm(1.25) if first_indent else Cm(0)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.space_after = Pt(0)
    return p

def add_image_placeholder(doc, num, caption):
    """Плейсхолдер для изображения"""
    # Пустая строка перед
    doc.add_paragraph()
    
    # Рамка-плейсхолдер
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'[ВСТАВИТЬ СКРИНШОТ #{num}]')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.italic = True
    p.paragraph_format.first_line_indent = Cm(0)
    
    # Подпись к рисунку
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'Рисунок {num} – {caption}')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_after = Pt(12)
    
    # Пустая строка после
    doc.add_paragraph()

def add_list_item(doc, text, level=0):
    """Элемент списка"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    indent = '    ' * level
    run = p.add_run(f'{indent}– {text}')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    p.paragraph_format.first_line_indent = Cm(1.25)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE


def add_code_block(doc, code, caption_num, caption_text):
    """Блок кода"""
    # Подпись перед кодом
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'Листинг {caption_num} – {caption_text}')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    p.paragraph_format.first_line_indent = Cm(0)
    
    # Код
    for line in code.split('\n'):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(line)
        run.font.name = 'Courier New'
        run.font.size = Pt(10)
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    
    doc.add_paragraph()

def add_toc(doc):
    """Оглавление"""
    add_heading1(doc, 'ОГЛАВЛЕНИЕ')
    
    toc_items = [
        ('Введение', 3),
        ('1. Описание программы', 5),
        ('1.1. Алгоритмические решения', 5),
        ('1.1.1. Безопасность', 5),
        ('1.1.2. Клиент', 8),
        ('1.1.3. Система бронирования', 10),
        ('1.2. Описание интерфейса программы', 11),
        ('1.2.1. Навигация и Футер', 11),
        ('1.2.2. Регистрация и Авторизация', 12),
        ('1.2.3. Главная страница', 14),
        ('1.2.4. Расписание рейсов', 15),
        ('1.2.5. Бронирование билетов', 16),
        ('1.2.6. Профиль пользователя', 17),
        ('1.2.7. Каталог воздушных судов', 18),
        ('1.2.8. Административная панель', 19),
        ('2. Архитектура приложения', 21),
        ('2.1. Зависимости проекта', 21),
        ('2.2. Клиент', 23),
        ('2.3. База данных', 24),
        ('3. Структура классов и их назначение в рамках проекта', 25),
        ('3.1. Сервер', 25),
        ('3.2. Клиент', 28),
        ('Заключение', 30),
        ('Список использованных источников', 31),
        ('Приложения', 33),
    ]
    
    for item, page in toc_items:
        p = doc.add_paragraph()
        # Определяем отступ по уровню
        if item.startswith('1.1.') or item.startswith('1.2.') or item.startswith('2.') or item.startswith('3.'):
            if '.' in item[2:] and item[0].isdigit():
                indent = Cm(1.25)
            else:
                indent = Cm(0.5)
        elif item[0].isdigit():
            indent = Cm(0)
        else:
            indent = Cm(0)
        
        run = p.add_run(item)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
        
        # Добавляем точки и номер страницы
        p.add_run('.' * 3 + str(page))
        p.paragraph_format.first_line_indent = indent
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    
    doc.add_page_break()

def add_introduction(doc):
    """Введение"""
    add_heading1(doc, 'ВВЕДЕНИЕ')
    
    add_paragraph_text(doc, 
        'Авиационная отрасль является одной из наиболее динамично развивающихся сфер мировой экономики. '
        'Ежегодно миллионы пассажиров пользуются услугами авиакомпаний, что создает необходимость в эффективных '
        'информационных системах для управления данными о рейсах, бронировании билетов и обслуживании клиентов [1]. '
        'Современные информационные технологии позволяют автоматизировать большинство процессов, связанных с '
        'деятельностью авиакомпаний, что существенно повышает качество обслуживания пассажиров и оптимизирует работу персонала [2].')
    
    add_paragraph_text(doc,
        'Информационные системы авиакомпаний представляют собой комплексные программные решения, обеспечивающие '
        'хранение, обработку и представление данных о рейсах, пассажирах, воздушных судах и бронированиях [3]. '
        'Такие системы должны обеспечивать высокую надежность, безопасность данных и удобство использования для '
        'различных категорий пользователей — от пассажиров до администраторов системы.')
    
    add_paragraph_text(doc,
        'Развитие веб-технологий и фреймворков для разработки корпоративных приложений, таких как Spring Boot, '
        'позволяет создавать масштабируемые и безопасные информационные системы с минимальными затратами времени '
        'и ресурсов [4]. Использование архитектурного паттерна MVC (Model-View-Controller) обеспечивает четкое '
        'разделение логики приложения, что упрощает его поддержку и развитие [5].')
    
    add_paragraph_text(doc,
        'Актуальность разработки информационной системы авиакомпании обусловлена необходимостью предоставления '
        'пользователям удобного инструмента для просмотра расписания рейсов, бронирования билетов и управления '
        'своими данными. Кроме того, система должна обеспечивать администраторам возможность эффективного '
        'управления данными о рейсах и воздушных судах.')
    
    add_paragraph_text(doc,
        'Объектом исследования является процесс информационного обеспечения деятельности авиакомпании.')
    
    add_paragraph_text(doc,
        'Предметом исследования являются методы и технологии разработки веб-приложений на основе фреймворка Spring Boot.')
    
    add_paragraph_text(doc,
        'Таким образом, целью проекта является разработка информационной системы авиакомпании, обеспечивающей '
        'автоматизацию процессов бронирования билетов и управления данными о рейсах.')
    
    add_paragraph_text(doc, 'Для достижения поставленной цели необходимо решить следующие задачи:')
    
    add_list_item(doc, 'разработка клиент-серверного приложения на языке программирования Java, предназначенного для информационного обеспечения деятельности авиакомпании;')
    add_list_item(doc, 'создание серверной части приложения на основе фреймворка Spring Boot, отвечающей за обработку запросов клиентов и предоставление данных о рейсах;')
    add_list_item(doc, 'разработка клиентской части приложения с графическим интерфейсом с помощью HTML, CSS и JavaScript, обеспечивающей удобный интерфейс для взаимодействия с сервером;')
    add_list_item(doc, 'применение модели MVC (Model View Controller) для разделения управляющей логики на отдельные компоненты, что позволяет улучшить фундаментальные свойства системы;')
    add_list_item(doc, 'реализация системы аутентификации и авторизации пользователей с разграничением прав доступа;')
    add_list_item(doc, 'создание функционала бронирования билетов с учетом доступности мест на рейсах.')
    
    add_paragraph_text(doc,
        'Для решения выше поставленных задач используются следующие инструментальные средства: Java 17, '
        'Spring Boot 3.3.4 [6-8], Spring Security, Spring Data JPA, Hibernate, MySQL, Thymeleaf и т.д.')
    
    add_paragraph_text(doc,
        'Таким образом, разработка информационной системы авиакомпании является актуальным инструментом для '
        'автоматизации процессов бронирования и управления данными в сфере авиаперевозок.')
    
    doc.add_page_break()


def add_section1(doc):
    """Раздел 1. Описание программы"""
    add_heading1(doc, '1. ОПИСАНИЕ ПРОГРАММЫ')
    
    add_heading2(doc, '1.1. Алгоритмические решения')
    
    add_heading3(doc, '1.1.1. Безопасность')
    
    add_paragraph_text(doc,
        'Отдельное внимание стоит уделить безопасности приложения и то, как она реализована (рисунок 1).')
    
    add_image_placeholder(doc, 1, 'Архитектура безопасности')
    
    add_paragraph_text(doc,
        'Для начала, после запуска проекта, настраивается цепочка фильтров безопасности (SecurityFilterChain), '
        'которая отвечает за основные моменты безопасности, например: какая страница будет отвечать за вход в '
        'приложение, на какие страницы будет доступ у пользователей, а на какие нужны будут отдельные разрешения, '
        'на какой адрес будет осуществляться выход пользователя из приложения.')
    
    add_paragraph_text(doc,
        'Затем инициализируется несколько бин-компонентов, которые отвечают за настройку AuthenticationManager. '
        'AuthenticationManager — это интерфейс в Spring Security, который отвечает за процесс аутентификации пользователей.')
    
    add_paragraph_text(doc,
        'AuthenticationManager принимает объект Authentication, который содержит учетные данные пользователя, '
        'и возвращает объект Authentication, который содержит информацию о пользователе, если аутентификация '
        'прошла успешно. Если аутентификация не удалась, AuthenticationManager может выбросить исключение '
        'AuthenticationException, которое указывает на причину ошибки.')
    
    add_paragraph_text(doc,
        'AuthenticationManager используется в Spring Security для проверки учетных данных, полученных от '
        'пользователя, и выполнения процесса аутентификации. Он настроен с помощью авто-конфигурации Spring Boot '
        'и работает с базой данных MySQL для хранения информации о пользователях.')
    
    add_paragraph_text(doc,
        'В Spring Security существует несколько реализаций интерфейса AuthenticationManager, таких как '
        'ProviderManager, которая делегирует аутентификацию другим провайдерам, и DaoAuthenticationProvider, '
        'который использует базу данных для аутентификации пользователей, что и используется в данном проекте.')
    
    add_paragraph_text(doc,
        'При настройке DaoAuthenticationProvider необходимо указать UserDetailsService и PasswordEncoder.')
    
    add_paragraph_text(doc,
        'UserDetailsService — это интерфейс в Spring Security, который отвечает за загрузку информации о '
        'пользователе из источника данных, такого как база данных.')
    
    add_paragraph_text(doc,
        'UserDetailsService содержит метод loadUserByUsername(), который принимает имя пользователя в качестве '
        'параметра и возвращает объект UserDetails, который содержит информацию о пользователе, такую как имя '
        'пользователя, пароль и список ролей. Если пользователь не найден, UserDetailsService выбрасывает '
        'исключение RuntimeException.')
    
    add_paragraph_text(doc,
        'PasswordEncoder отвечает за хэширование паролей пользователей для обеспечения безопасности в приложениях.')
    
    add_paragraph_text(doc,
        'PasswordEncoder принимает пароль пользователя в качестве параметра и возвращает хэш-код, который '
        'сохраняется в базе данных. При аутентификации пользователя, введенный пользователем пароль сравнивается '
        'с сохраненным хэш-кодом. В программе используется алгоритм шифрования BCrypt.')
    
    add_paragraph_text(doc,
        'Таким образом, AuthenticationManager — это интерфейс в Spring Security, который отвечает за процесс '
        'аутентификации пользователей. Он используется для проверки учетных данных пользователя.')
    
    add_paragraph_text(doc,
        'В программе реализована защита от межсайтовой подделки запросов CSRF (Cross-Site Request Forgery) '
        'путем отключения данной защиты для упрощения разработки, однако в продакшн-версии рекомендуется её включить.')
    
    add_paragraph_text(doc,
        'При регистрации по умолчанию устанавливается роль обычного пользователя (USER). Для получения '
        'дополнительных прав (ADMIN, MANAGER) необходимо обратиться к администратору системы.')
    
    add_paragraph_text(doc, 'В системе реализовано три роли пользователей:')
    add_list_item(doc, 'USER — обычный пользователь, имеющий доступ к просмотру расписания, бронированию билетов и управлению своим профилем;')
    add_list_item(doc, 'MANAGER — менеджер, имеющий дополнительные права на редактирование контента страниц;')
    add_list_item(doc, 'ADMIN — администратор, имеющий полный доступ ко всем функциям системы, включая управление рейсами и воздушными судами.')

def add_section1_2(doc):
    """Раздел 1.1.2. Клиент"""
    add_heading3(doc, '1.1.2. Клиент')
    
    add_paragraph_text(doc,
        'В приложении используется архитектура MVC (Model-View-Controller), которая разделяет приложение на '
        'три основных компонента: модель, представление и контроллер. Взаимодействие всех элементов можно '
        'рассмотреть на рисунке 2.')
    
    add_image_placeholder(doc, 2, 'Взаимодействие в паттерне MVC')
    
    add_paragraph_text(doc,
        'Модель (Model) — это компонент, отвечающий за обработку данных. Он содержит классы, которые представляют '
        'данные и методы для работы с ними. Например, класс User, который представляет информацию о пользователе, '
        'класс Flight, представляющий данные о рейсе, класс Aircraft для информации о воздушных судах, и класс '
        'Booking для хранения данных о бронированиях.')
    
    add_paragraph_text(doc,
        'Представление (View) — это компонент, который отображает данные в пользовательском интерфейсе. Он '
        'содержит файлы представлений, например HTML-страницы, которые отображают данные, полученные из модели. '
        'В Spring Boot для рендеринга представлений используется шаблонизатор Thymeleaf.')
    
    add_paragraph_text(doc,
        'Контроллер (Controller) — это компонент, который связывает модель и представление. Он содержит методы, '
        'которые обрабатывают запросы от пользователя и возвращают данные в нужном формате. Например, '
        'AuthController, который содержит методы для обработки запросов, связанных с аутентификацией пользователей, '
        'BookingController для обработки бронирований, ScheduleController для отображения расписания рейсов.')
    
    add_paragraph_text(doc, 'Взаимодействие пользователя и приложения:')
    add_list_item(doc, 'пользователь открывает браузер и вводит URL-адрес приложения;')
    add_list_item(doc, 'сервер принимает запрос от пользователя и сопоставляет URL-адрес с соответствующим методом контроллера;')
    add_list_item(doc, 'метод контроллера обращается к методам JPA-репозитория для извлечения данных из БД, их обработки и последующей передачи в представление;')
    add_list_item(doc, 'представление использует данные, полученные из контроллера, и отображает их в пользовательском интерфейсе с помощью шаблонизатора Thymeleaf;')
    add_list_item(doc, 'пользователь взаимодействует с интерфейсом, например, заполняет форму бронирования и отправляет её;')
    add_list_item(doc, 'сервер принимает запрос от пользователя, сопоставляет URL-адрес с соответствующим методом контроллера и обрабатывает запрос;')
    add_list_item(doc, 'метод контроллера обрабатывает запрос от пользователя, выполняет необходимые операции с базой данных и возвращает результат;')
    add_list_item(doc, 'представление использует данные, переданные из контроллера, и отображает их в пользовательском интерфейсе;')
    add_list_item(doc, 'пользователь продолжает взаимодействовать с интерфейсом, и процесс повторяется.')
    
    add_paragraph_text(doc,
        'Таким образом, пользователь взаимодействует с приложением через браузер, а архитектура MVC обеспечивает '
        'четкое разделение ответственности между компонентами системы.')


def add_section1_3(doc):
    """Раздел 1.1.3. Система бронирования"""
    add_heading3(doc, '1.1.3. Система бронирования')
    
    add_paragraph_text(doc,
        'Система бронирования билетов является ключевым функциональным модулем информационной системы авиакомпании. '
        'Она обеспечивает возможность пользователям бронировать места на рейсы с учетом доступности мест.')
    
    add_paragraph_text(doc, 'Основные компоненты системы бронирования включают:')
    add_list_item(doc, 'Проверка аутентификации — перед бронированием система проверяет, авторизован ли пользователь. Неавторизованные пользователи перенаправляются на страницу входа;')
    add_list_item(doc, 'Проверка роли пользователя — администраторы не могут бронировать билеты, так как эта функция предназначена только для обычных пользователей;')
    add_list_item(doc, 'Проверка доступности мест — система проверяет количество доступных мест на рейсе. Если места закончились, пользователь получает соответствующее уведомление;')
    add_list_item(doc, 'Проверка дублирования бронирования — система не позволяет одному пользователю забронировать один и тот же рейс дважды;')
    add_list_item(doc, 'Создание бронирования — при успешном прохождении всех проверок создается запись о бронировании с указанием паспортных данных и адреса пользователя;')
    add_list_item(doc, 'Обновление количества мест — после успешного бронирования количество доступных мест на рейсе уменьшается на единицу.')
    
    add_paragraph_text(doc,
        'Процесс бронирования реализован с использованием транзакций (@Transactional), что обеспечивает целостность '
        'данных при одновременном доступе нескольких пользователей к системе.')
    
    add_paragraph_text(doc,
        'Пользователь также имеет возможность отменить бронирование через свой профиль. При отмене бронирования '
        'количество доступных мест на рейсе увеличивается на единицу, а запись о бронировании удаляется из базы данных.')

def add_section1_interface(doc):
    """Раздел 1.2. Описание интерфейса программы"""
    add_heading2(doc, '1.2. Описание интерфейса программы')
    
    add_paragraph_text(doc,
        'В данном разделе описывается интерфейс программы, написанный на языках HTML, CSS и JavaScript с '
        'использованием шаблонизатора Thymeleaf.')
    
    add_paragraph_text(doc,
        'Основной целью интерфейса является предоставление пользователю удобного и интуитивно понятного способа '
        'взаимодействия с программой. Для этого он включает в себя различные элементы управления, такие как кнопки, '
        'текстовые поля, выпадающие списки, таблицы и формы, которые позволяют пользователю выполнить необходимые действия.')
    
    # 1.2.1. Навигация и Футер
    add_heading3(doc, '1.2.1. Навигация и Футер')
    
    add_paragraph_text(doc,
        'В интерфейсе реализована удобная навигация в верхней части экрана (рисунок 3).')
    
    add_image_placeholder(doc, 3, 'Навигация сайта')
    
    add_paragraph_text(doc, 'Навигация в браузере выглядит следующим образом:')
    add_list_item(doc, 'Главная — переход на главную страницу системы;')
    add_list_item(doc, 'Расписание — просмотр расписания всех рейсов;')
    add_list_item(doc, 'Билеты — страница поиска и бронирования билетов;')
    add_list_item(doc, 'Воздушные суда — каталог воздушных судов авиакомпании;')
    add_list_item(doc, 'Радар — информационная страница о радаре;')
    add_list_item(doc, 'О нас — информация о системе и авиакомпании.')
    
    add_paragraph_text(doc,
        'Если пользователь не авторизован, то также отображаются ссылки на "Войти" и "Регистрацию". Если '
        'пользователь авторизован, то отображается ссылка на профиль и кнопка "Выйти", которая позволяет '
        'пользователю выйти из своей учетной записи.')
    
    add_paragraph_text(doc,
        'Для администраторов в навигации появляются дополнительные пункты меню для управления рейсами и воздушными судами.')
    
    add_paragraph_text(doc,
        'Футер (англ. footer) — это нижняя часть веб-страницы, которая содержит дополнительную информацию и '
        'ссылки на другие страницы или ресурсы (рисунок 4).')
    
    add_image_placeholder(doc, 4, 'Футер')
    
    add_paragraph_text(doc, 'В данном проекте футер содержит следующие элементы:')
    add_list_item(doc, 'Контактная информация авиакомпании;')
    add_list_item(doc, 'Ссылки на основные разделы сайта;')
    add_list_item(doc, 'Информация об авторских правах.')
    
    # 1.2.2. Регистрация и Авторизация
    add_heading3(doc, '1.2.2. Регистрация и Авторизация')
    
    add_paragraph_text(doc,
        'Регистрация пользователя — это процесс, который позволяет создать учетную запись на сайте и получить '
        'доступ к его функциональности. В данном проекте, для регистрации пользователя ему необходимо внести '
        'свой логин, пароль, полное имя, электронную почту и номер телефона (рисунок 5).')
    
    add_image_placeholder(doc, 5, 'Регистрация')
    
    add_paragraph_text(doc, 'При регистрации пользователь должен заполнить все требуемые поля:')
    add_list_item(doc, 'Логин (username) — уникальное имя пользователя для входа в систему;')
    add_list_item(doc, 'Пароль — пароль, который будет использоваться для входа на сайт;')
    add_list_item(doc, 'Полное имя (fullName) — ФИО пользователя;')
    add_list_item(doc, 'Email — адрес электронной почты;')
    add_list_item(doc, 'Телефон — контактный номер телефона.')
    
    add_paragraph_text(doc,
        'После заполнения всех полей пользователь должен нажать на кнопку "Зарегистрироваться". При успешной '
        'регистрации система перенаправляет пользователя на страницу входа в приложение с сообщением об успешной '
        'регистрации (рисунок 6).')
    
    add_image_placeholder(doc, 6, 'Авторизация')
    
    add_paragraph_text(doc, 'При регистрации пользователю следует учитывать следующее:')
    add_list_item(doc, 'Логин должен быть уникальным — если пользователь с таким логином уже существует, система выведет сообщение об ошибке;')
    add_list_item(doc, 'Все обязательные поля должны быть заполнены;')
    add_list_item(doc, 'Email должен быть уникальным в системе.')
    
    add_paragraph_text(doc,
        'Авторизация — это процесс, который позволяет пользователю войти в свою учетную запись на сайте, '
        'используя свои учетные данные.')
    
    add_paragraph_text(doc,
        'В данном проекте, для авторизации пользователю необходимо ввести только свой логин и пароль. После '
        'ввода данных пользователь должен нажать на кнопку "Войти", после чего система проверит правильность '
        'введенных данных и, если они корректны, предоставит пользователю доступ к его учетной записи. Если '
        'данные оказались неверны, пользователя вернет на страницу авторизации с сообщением об ошибке.')


def add_section1_interface_continued(doc):
    """Продолжение раздела 1.2"""
    # 1.2.3. Главная страница
    add_heading3(doc, '1.2.3. Главная страница')
    
    add_paragraph_text(doc,
        'Главная страница сайта — это первая страница, которую пользователь увидит при заходе на сайт. Её '
        'основная задача — представить проект и привлечь внимание пользователя, а также предоставить ему '
        'информацию о доступной функциональности и возможностях системы (рисунок 7).')
    
    add_image_placeholder(doc, 7, 'Главная страница')
    
    add_paragraph_text(doc, 'В данном проекте главная страница содержит следующие элементы:')
    add_list_item(doc, 'Шапка страницы — верхняя часть страницы с логотипом проекта, названием и основным меню навигации;')
    add_list_item(doc, 'Приветственный блок — содержит название системы и краткое описание её назначения;')
    add_list_item(doc, 'Видео-фон — для создания визуально привлекательного интерфейса используется фоновое видео;')
    add_list_item(doc, 'Быстрые ссылки — блоки с переходами на основные разделы системы;')
    add_list_item(doc, 'Футер — нижняя часть страницы с дополнительной информацией.')
    
    add_paragraph_text(doc,
        'Контент главной страницы может редактироваться администраторами и менеджерами через специальную '
        'административную панель.')
    
    # 1.2.4. Расписание рейсов
    add_heading3(doc, '1.2.4. Расписание рейсов')
    
    add_paragraph_text(doc,
        'Страница расписания рейсов представляет собой интерфейс, на котором пользователь может просмотреть '
        'информацию обо всех рейсах авиакомпании (рисунок 8).')
    
    add_image_placeholder(doc, 8, 'Расписание рейсов')
    
    add_paragraph_text(doc, 'На странице расписания отображается таблица со следующей информацией о каждом рейсе:')
    add_list_item(doc, 'Номер рейса (flightNumber);')
    add_list_item(doc, 'Авиакомпания (airline);')
    add_list_item(doc, 'Пункт отправления (origin);')
    add_list_item(doc, 'Пункт назначения (destination);')
    add_list_item(doc, 'Время вылета (departureTime);')
    add_list_item(doc, 'Время прилета (arrivalTime);')
    add_list_item(doc, 'Выход (gate);')
    add_list_item(doc, 'Тип воздушного судна (aircraftType);')
    add_list_item(doc, 'Статус рейса (status);')
    add_list_item(doc, 'Количество доступных мест (availableSeats).')
    
    add_paragraph_text(doc, 'Для администраторов на странице расписания доступны дополнительные функции:')
    add_list_item(doc, 'Добавление нового рейса;')
    add_list_item(doc, 'Редактирование существующего рейса;')
    add_list_item(doc, 'Удаление рейса.')
    
    # 1.2.5. Бронирование билетов
    add_heading3(doc, '1.2.5. Бронирование билетов')
    
    add_paragraph_text(doc,
        'Страница бронирования билетов представляет собой интерфейс, на котором пользователь может найти '
        'подходящий рейс и забронировать билет (рисунок 9).')
    
    add_image_placeholder(doc, 9, 'Страница бронирования билетов')
    
    add_paragraph_text(doc, 'На странице реализованы следующие функции:')
    add_list_item(doc, 'Поиск рейсов — пользователь может ввести пункт отправления и пункт назначения для фильтрации рейсов;')
    add_list_item(doc, 'Сортировка — результаты можно отсортировать по времени вылета, прилета, номеру рейса, авиакомпании или статусу;')
    add_list_item(doc, 'Бронирование — для каждого рейса с доступными местами отображается кнопка "Забронировать".')
    
    add_paragraph_text(doc,
        'При нажатии на кнопку "Забронировать" пользователь переходит на страницу подтверждения бронирования (рисунок 10).')
    
    add_image_placeholder(doc, 10, 'Подтверждение бронирования')
    
    add_paragraph_text(doc, 'На странице подтверждения пользователь должен ввести:')
    add_list_item(doc, 'Номер паспорта;')
    add_list_item(doc, 'Адрес проживания.')
    
    add_paragraph_text(doc,
        'После заполнения данных и нажатия кнопки "Подтвердить бронирование" создается запись о бронировании, '
        'и пользователь перенаправляется в свой профиль.')
    
    # 1.2.6. Профиль пользователя
    add_heading3(doc, '1.2.6. Профиль пользователя')
    
    add_paragraph_text(doc,
        'Страница профиля пользователя предоставляет возможность просмотра и редактирования личной информации, '
        'а также управления бронированиями (рисунок 11).')
    
    add_image_placeholder(doc, 11, 'Профиль пользователя')
    
    add_paragraph_text(doc, 'На странице профиля отображается следующая информация:')
    add_list_item(doc, 'Логин пользователя;')
    add_list_item(doc, 'Полное имя;')
    add_list_item(doc, 'Адрес электронной почты;')
    add_list_item(doc, 'Номер телефона;')
    add_list_item(doc, 'Роль в системе.')
    
    add_paragraph_text(doc,
        'Также на странице профиля отображается список всех бронирований пользователя с информацией о номере '
        'рейса, маршруте, дате и времени вылета, дате бронирования.')
    
    add_paragraph_text(doc,
        'Для каждого бронирования доступна кнопка "Отменить", которая позволяет отменить бронирование. При '
        'отмене бронирования количество доступных мест на рейсе увеличивается.')
    
    add_paragraph_text(doc,
        'Пользователь может редактировать свой профиль, нажав на кнопку "Редактировать профиль" (рисунок 12).')
    
    add_image_placeholder(doc, 12, 'Редактирование профиля')
    
    add_paragraph_text(doc, 'При редактировании профиля пользователь может изменить:')
    add_list_item(doc, 'Полное имя;')
    add_list_item(doc, 'Адрес электронной почты;')
    add_list_item(doc, 'Номер телефона.')
    
    add_paragraph_text(doc, 'Логин и роль пользователя изменить нельзя.')
    
    # 1.2.7. Каталог воздушных судов
    add_heading3(doc, '1.2.7. Каталог воздушных судов')
    
    add_paragraph_text(doc,
        'Страница каталога воздушных судов представляет собой галерею с информацией о воздушных судах '
        'авиакомпании (рисунок 13).')
    
    add_image_placeholder(doc, 13, 'Каталог воздушных судов')
    
    add_paragraph_text(doc, 'На странице отображаются карточки воздушных судов, каждая из которых содержит:')
    add_list_item(doc, 'Изображение воздушного судна;')
    add_list_item(doc, 'Название модели;')
    add_list_item(doc, 'Краткое описание.')
    
    add_paragraph_text(doc,
        'При нажатии на карточку пользователь переходит на страницу с подробной информацией о воздушном судне (рисунок 14).')
    
    add_image_placeholder(doc, 14, 'Детальная информация о воздушном судне')
    
    add_paragraph_text(doc, 'На странице детальной информации отображается:')
    add_list_item(doc, 'Полноразмерное изображение воздушного судна;')
    add_list_item(doc, 'Название модели;')
    add_list_item(doc, 'Подробное описание технических характеристик и особенностей.')
    
    # 1.2.8. Административная панель
    add_heading3(doc, '1.2.8. Административная панель')
    
    add_paragraph_text(doc,
        'Административная панель доступна только пользователям с ролью ADMIN и предоставляет функции '
        'управления данными системы (рисунок 15).')
    
    add_image_placeholder(doc, 15, 'Административная панель')
    
    add_paragraph_text(doc, 'Администратор имеет доступ к следующим функциям:')
    add_list_item(doc, 'Управление рейсами: добавление, редактирование, удаление рейсов;')
    add_list_item(doc, 'Управление воздушными судами: добавление с загрузкой изображения, редактирование, удаление;')
    add_list_item(doc, 'Управление контентом страниц: редактирование текста главной страницы, страницы "О нас", страницы "Радар".')
    
    add_paragraph_text(doc, 'Форма добавления/редактирования рейса (рисунок 16) содержит следующие поля:')
    
    add_image_placeholder(doc, 16, 'Форма редактирования рейса')
    
    add_list_item(doc, 'Номер рейса;')
    add_list_item(doc, 'Авиакомпания;')
    add_list_item(doc, 'Пункт отправления;')
    add_list_item(doc, 'Пункт назначения;')
    add_list_item(doc, 'Время вылета;')
    add_list_item(doc, 'Время прилета;')
    add_list_item(doc, 'Выход;')
    add_list_item(doc, 'Тип воздушного судна;')
    add_list_item(doc, 'Статус;')
    add_list_item(doc, 'Количество доступных мест.')
    
    add_paragraph_text(doc, 'Форма добавления/редактирования воздушного судна (рисунок 17) содержит:')
    
    add_image_placeholder(doc, 17, 'Форма редактирования воздушного судна')
    
    add_list_item(doc, 'Название модели;')
    add_list_item(doc, 'Изображение (загрузка файла);')
    add_list_item(doc, 'Краткое описание;')
    add_list_item(doc, 'Подробное описание.')
    
    doc.add_page_break()


def add_section2(doc):
    """Раздел 2. Архитектура приложения"""
    add_heading1(doc, '2. АРХИТЕКТУРА ПРИЛОЖЕНИЯ')
    
    add_heading2(doc, '2.1. Зависимости проекта')
    
    add_paragraph_text(doc,
        'В файле pom.xml представлены зависимости проекта, которые определяют, какие библиотеки и фреймворки '
        'будут использоваться в проекте. Ниже приведен список всех зависимостей, которые указаны в данном файле:')
    
    add_list_item(doc, 'spring-boot-starter-web — зависимость для создания веб-приложения на Spring. Она включает в себя все необходимые библиотеки и компоненты, такие как Tomcat, Jackson, Spring MVC;')
    add_list_item(doc, 'spring-boot-starter-thymeleaf — зависимость для работы с шаблонизатором Thymeleaf, который используется для рендеринга HTML-страниц на стороне сервера;')
    add_list_item(doc, 'thymeleaf-extras-springsecurity6 — зависимость для интеграции Thymeleaf с Spring Security, позволяющая использовать директивы безопасности в шаблонах;')
    add_list_item(doc, 'spring-boot-starter-data-jpa — зависимость для работы с базой данных с использованием JPA. Она включает Hibernate, Spring Data JPA;')
    add_list_item(doc, 'spring-boot-starter-security — зависимость для создания безопасных веб-приложений с аутентификацией и авторизацией;')
    add_list_item(doc, 'mysql-connector-j — зависимость для работы с базой данных MySQL, предоставляющая JDBC-драйвер;')
    add_list_item(doc, 'spring-boot-starter-test — зависимость для написания тестов, включающая JUnit, Mockito.')
    
    add_paragraph_text(doc, 'Структура файла pom.xml следующая (рисунок 18):')
    
    add_image_placeholder(doc, 18, 'Фрагмент кода зависимостей')
    
    add_list_item(doc, 'project — корневой элемент файла, определяющий проект с информацией о версии, артефакте, группе и имени;')
    add_list_item(doc, 'modelVersion — элемент, указывающий на версию модели проекта (4.0.0);')
    add_list_item(doc, 'properties — элемент с пользовательскими свойствами, в данном случае версией Java (17);')
    add_list_item(doc, 'dependencies — элемент, определяющий зависимости проекта от других библиотек;')
    add_list_item(doc, 'build — элемент конфигурации сборки проекта, включая плагин spring-boot-maven-plugin.')
    
    add_heading2(doc, '2.2. Клиент')
    
    add_paragraph_text(doc, 'На стороне клиента используется HTML, CSS, JavaScript и Thymeleaf.')
    
    add_paragraph_text(doc,
        'HTML является основой веб-страниц и используется для создания структуры и содержимого интерфейса. '
        'CSS используется для стилизации веб-страниц и задания внешнего вида элементов интерфейса. В проекте '
        'используется собственный файл стилей app.css.')
    
    add_paragraph_text(doc,
        'JavaScript используется для создания интерактивности и динамических элементов интерфейса. Thymeleaf — '
        'это шаблонизатор для веб-приложений, который позволяет создавать HTML-шаблоны с использованием Java-кода '
        'и вставлять данные из модели Java в HTML-шаблоны.')
    
    add_paragraph_text(doc, 'Thymeleaf предоставляет следующие возможности:')
    add_list_item(doc, 'Вставка данных из модели в HTML с помощью атрибутов th:text, th:value;')
    add_list_item(doc, 'Условный рендеринг с помощью th:if, th:unless;')
    add_list_item(doc, 'Итерация по коллекциям с помощью th:each;')
    add_list_item(doc, 'Формирование URL-адресов с помощью th:href, th:action;')
    add_list_item(doc, 'Интеграция с Spring Security для отображения элементов в зависимости от роли пользователя (sec:authorize).')
    
    add_paragraph_text(doc,
        'Таким образом, использование HTML, CSS, JavaScript и Thymeleaf позволяет создавать красивый, '
        'адаптивный и интерактивный веб-интерфейс для проекта.')
    
    add_heading2(doc, '2.3. База данных')
    
    add_paragraph_text(doc,
        'Настройки базы данных находятся в файле application.properties (рисунок 19) в проекте. В них '
        'определяются параметры подключения к базе данных, а также настройки JPA.')
    
    add_image_placeholder(doc, 19, 'Файл настроек подключения')
    
    add_paragraph_text(doc,
        'Данный код содержит настройки для подключения к базе данных MySQL и использования JPA (Java Persistence API) '
        'для работы с базой данных.')
    
    add_paragraph_text(doc, 'В блоке spring.datasource задаются параметры подключения к базе данных:')
    add_list_item(doc, 'url — URL-адрес базы данных (jdbc:mysql://localhost:3306/airline_db) с параметрами автоматического создания базы данных;')
    add_list_item(doc, 'username — имя пользователя базы данных;')
    add_list_item(doc, 'password — пароль пользователя базы данных;')
    add_list_item(doc, 'driver-class-name — имя драйвера (com.mysql.cj.jdbc.Driver).')
    
    add_paragraph_text(doc, 'В блоке spring.jpa задаются настройки JPA:')
    add_list_item(doc, 'database-platform — диалект Hibernate для работы с MySQL (org.hibernate.dialect.MySQLDialect);')
    add_list_item(doc, 'hibernate.ddl-auto — режим автоматического создания/обновления схемы базы данных (update);')
    add_list_item(doc, 'show-sql — включение вывода SQL-запросов в консоль для отладки;')
    add_list_item(doc, 'properties.hibernate.format_sql — форматирование SQL-запросов для удобства чтения.')
    
    add_paragraph_text(doc, 'Дополнительные настройки:')
    add_list_item(doc, 'spring.web.resources.static-locations — указание расположения статических ресурсов, включая папку uploads;')
    add_list_item(doc, 'spring.servlet.multipart.max-file-size — максимальный размер загружаемого файла (10MB);')
    add_list_item(doc, 'spring.servlet.multipart.max-request-size — максимальный размер запроса (10MB).')
    
    add_paragraph_text(doc,
        'Таким образом, эти настройки позволяют приложению подключаться к базе данных MySQL и использовать JPA '
        'для работы с данными в базе данных.')
    
    doc.add_page_break()


def add_section3(doc):
    """Раздел 3. Структура классов"""
    add_heading1(doc, '3. СТРУКТУРА КЛАССОВ И ИХ НАЗНАЧЕНИЕ В РАМКАХ ПРОЕКТА')
    
    add_heading2(doc, '3.1. Сервер')
    
    add_paragraph_text(doc,
        'На стороне сервера написано 22 класса, которые разбиты на структуру пакетов (рисунок 20).')
    
    add_image_placeholder(doc, 20, 'Структура классов')
    
    add_paragraph_text(doc,
        'Структура программы создана в удобном формате, где каждый пакет отвечает только за свою логику и свои классы.')
    
    add_paragraph_text(doc,
        'config — пакет, содержащий классы конфигурации приложения, которые определяют различные настройки и '
        'параметры, связанные с работой приложения:')
    add_list_item(doc, 'SecurityConfig — класс конфигурации для Spring Security, который определяет настройки безопасности приложения. Он определяет бины для объектов SecurityFilterChain, DaoAuthenticationProvider, UserDetailsService и PasswordEncoder;')
    add_list_item(doc, 'WebConfig — класс конфигурации для Spring MVC, который определяет обработчики ресурсов и другие настройки веб-приложения.')
    
    add_paragraph_text(doc,
        'controller — пакет, содержащий классы контроллеров, которые обрабатывают запросы от клиента и управляют '
        'логикой обработки этих запросов:')
    add_list_item(doc, 'AuthController — контроллер для авторизации и регистрации пользователей;')
    add_list_item(doc, 'MainPagesController — контроллер для главных страниц приложения (главная, о нас, радар);')
    add_list_item(doc, 'ScheduleController — контроллер для отображения расписания рейсов;')
    add_list_item(doc, 'TicketsController — контроллер для страницы бронирования билетов с функциями поиска и сортировки;')
    add_list_item(doc, 'BookingController — контроллер для бронирования билетов с проверкой доступности мест;')
    add_list_item(doc, 'ProfileController — контроллер для страницы профиля пользователя;')
    add_list_item(doc, 'AircraftController — контроллер для отображения каталога воздушных судов;')
    add_list_item(doc, 'AdminAircraftController — контроллер для административного управления воздушными судами;')
    add_list_item(doc, 'AdminFlightController — контроллер для административного управления рейсами;')
    add_list_item(doc, 'ManagerContentController — контроллер для управления контентом страниц;')
    add_list_item(doc, 'CustomErrorController — контроллер для обработки ошибок.')
    
    add_paragraph_text(doc,
        'handler — пакет, содержащий классы обработчиков событий:')
    add_list_item(doc, 'CustomAuthenticationSuccessHandler — класс, который обрабатывает успешную аутентификацию пользователя.')
    
    add_paragraph_text(doc,
        'initializer — пакет, содержащий классы для инициализации данных:')
    add_list_item(doc, 'DataInitializer — класс, который инициализирует начальные данные в базе данных при запуске приложения.')
    
    add_paragraph_text(doc,
        'model — пакет, содержащий классы моделей, которые представляют сущности и объекты данных в приложении:')
    add_list_item(doc, 'User — класс, представляющий информацию о пользователе системы с атрибутами id, username, password, fullName, email, phone и role;')
    add_list_item(doc, 'Role — перечисление (enum), определяющее возможные роли пользователей: USER, ADMIN, MANAGER;')
    add_list_item(doc, 'Flight — класс, представляющий данные о рейсе с атрибутами id, flightNumber, airline, origin, destination, departureTime, arrivalTime, gate, aircraftType, status и availableSeats;')
    add_list_item(doc, 'Aircraft — класс, представляющий данные о воздушном судне с атрибутами id, name, imageUrl, shortInfo и details;')
    add_list_item(doc, 'Booking — класс, представляющий данные о бронировании с атрибутами id, user, flight, bookedAt, passportNumber и address;')
    add_list_item(doc, 'PageContent — класс для хранения контента страниц с атрибутами id, title, content и body.')
    
    add_paragraph_text(doc,
        'repository — пакет, содержащий интерфейсы репозиториев, которые предоставляют доступ к данным в базе данных:')
    add_list_item(doc, 'UserRepository — репозиторий для управления данными о пользователях с методом findByUsername;')
    add_list_item(doc, 'FlightRepository — репозиторий для управления данными о рейсах с методом поиска по пунктам отправления и назначения;')
    add_list_item(doc, 'AircraftRepository — репозиторий для управления данными о воздушных судах;')
    add_list_item(doc, 'BookingRepository — репозиторий для управления данными о бронированиях с методами findByUser, existsByUserAndFlight и deleteByFlightId;')
    add_list_item(doc, 'PageContentRepository — репозиторий для управления контентом страниц.')
    
    add_heading2(doc, '3.2. Клиент')
    
    add_paragraph_text(doc,
        'Всего на стороне клиента используется 18 HTML-файлов шаблонов, 1 файл CSS для стилизации, изображения '
        'и видео для оформления интерфейса (рисунок 21).')
    
    add_image_placeholder(doc, 21, 'Структура клиентской части')
    
    add_paragraph_text(doc, 'Templates (шаблоны Thymeleaf):')
    add_list_item(doc, 'index.html — главная страница приложения с приветственным блоком и видео-фоном;')
    add_list_item(doc, 'login.html — страница авторизации пользователя;')
    add_list_item(doc, 'register.html — страница регистрации нового пользователя;')
    add_list_item(doc, 'schedule.html — страница расписания рейсов с таблицей всех рейсов;')
    add_list_item(doc, 'tickets.html — страница бронирования билетов с формой поиска и сортировкой;')
    add_list_item(doc, 'booking-confirm.html — страница подтверждения бронирования;')
    add_list_item(doc, 'profile.html — страница профиля пользователя;')
    add_list_item(doc, 'profile-edit.html — страница редактирования профиля;')
    add_list_item(doc, 'aircraft-list.html — страница каталога воздушных судов;')
    add_list_item(doc, 'aircraft-detail.html — страница детальной информации о воздушном судне;')
    add_list_item(doc, 'aircraft-form.html — форма добавления/редактирования воздушного судна;')
    add_list_item(doc, 'flight-form.html — форма добавления/редактирования рейса;')
    add_list_item(doc, 'about.html — страница "О нас";')
    add_list_item(doc, 'radar.html — информационная страница о радаре;')
    add_list_item(doc, 'manage-home-edit.html, manage-about-edit.html, manage-radar-edit.html — страницы редактирования контента;')
    add_list_item(doc, 'error.html — страница отображения ошибок.')
    
    add_paragraph_text(doc, 'Static (статические ресурсы):')
    add_list_item(doc, 'css/app.css — основной файл стилей приложения;')
    add_list_item(doc, 'images/ — папка с изображениями (логотип, логотип университета);')
    add_list_item(doc, 'videos/ — папка с видео-файлами для фона главной страницы.')
    
    add_paragraph_text(doc, 'Uploads (загруженные файлы):')
    add_list_item(doc, 'uploads/aircraft/ — папка для хранения загруженных изображений воздушных судов;')
    add_list_item(doc, 'uploads/staff/ — папка для хранения изображений персонала.')
    
    doc.add_page_break()


def add_conclusion(doc):
    """Заключение"""
    add_heading1(doc, 'ЗАКЛЮЧЕНИЕ')
    
    add_paragraph_text(doc,
        'В результате выполнения курсовой работы была разработана информационная система авиакомпании с '
        'использованием фреймворка Spring Boot и создано соответствующее веб-приложение. Основной целью проекта '
        'было разработать систему, позволяющую автоматизировать процессы бронирования билетов и управления '
        'данными о рейсах, а также предоставляющую пользователям удобный интерфейс для взаимодействия с системой.')
    
    add_paragraph_text(doc, 'Для достижения этой цели были выполнены следующие задачи:')
    add_list_item(doc, 'разработка клиент-серверного приложения на языке Java с использованием Spring Boot 3.3.4;')
    add_list_item(doc, 'создание серверной части приложения с реализацией контроллеров для обработки запросов;')
    add_list_item(doc, 'разработка клиентской части приложения с графическим интерфейсом с помощью HTML, CSS, JavaScript и шаблонизатора Thymeleaf;')
    add_list_item(doc, 'применение модели MVC для разделения управляющей логики на отдельные компоненты;')
    add_list_item(doc, 'реализация системы аутентификации и авторизации с использованием Spring Security и разграничением прав доступа для трех ролей (USER, ADMIN, MANAGER);')
    add_list_item(doc, 'создание функционала бронирования билетов с учетом доступности мест и предотвращением дублирования бронирований;')
    add_list_item(doc, 'реализация административной панели для управления рейсами, воздушными судами и контентом страниц.')
    
    add_paragraph_text(doc, 'В результате работы была получена эффективная и удобная система, позволяющая:')
    add_list_item(doc, 'пользователям просматривать расписание рейсов, искать и бронировать билеты, управлять своим профилем и бронированиями;')
    add_list_item(doc, 'администраторам управлять данными о рейсах и воздушных судах, редактировать контент страниц;')
    add_list_item(doc, 'менеджерам редактировать контент информационных страниц.')
    
    add_paragraph_text(doc,
        'Разработанное приложение может быть использовано в качестве основы для создания полноценной '
        'информационной системы авиакомпании с возможностью дальнейшего расширения функциональности.')
    
    add_paragraph_text(doc,
        'Также, во время выполнения курсовой работы были получены навыки разработки веб-приложений на основе '
        'фреймворка Spring Boot, работы с базами данных через JPA/Hibernate, реализации системы безопасности '
        'с использованием Spring Security. Эти навыки могут быть применены в различных областях, где требуется '
        'разработка корпоративных веб-приложений.')
    
    doc.add_page_break()

def add_references(doc):
    """Список использованных источников"""
    add_heading1(doc, 'СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ')
    
    references = [
        'Беляев, С.А. Информационные системы в авиации: учебное пособие / С.А. Беляев. – М.: Транспорт, 2020. – 256 с.',
        'Гагарина, Л.Г. Технология разработки программного обеспечения: учебное пособие / Л.Г. Гагарина, Е.В. Кокорева, Б.Д. Виснадул. – М.: ФОРУМ: ИНФРА-М, 2019. – 400 с.',
        'Орлов, С.А. Технологии разработки программного обеспечения: учебник / С.А. Орлов. – СПб.: Питер, 2021. – 608 с.',
        'Walls, C. Spring в действии / C. Walls. – М.: ДМК Пресс, 2022. – 544 с.',
        'Gamma, E. Design Patterns: Elements of Reusable Object-Oriented Software / E. Gamma, R. Helm, R. Johnson, J. Vlissides. – Addison-Wesley, 1994. – 416 p.',
        'Spring Boot Reference Documentation [Электронный ресурс]. – Режим доступа: https://docs.spring.io/spring-boot/docs/current/reference/htmlsingle/ (дата обращения: 15.12.2024).',
        'Spring Security Reference [Электронный ресурс]. – Режим доступа: https://docs.spring.io/spring-security/reference/index.html (дата обращения: 15.12.2024).',
        'Spring Data JPA Reference Documentation [Электронный ресурс]. – Режим доступа: https://docs.spring.io/spring-data/jpa/docs/current/reference/html/ (дата обращения: 15.12.2024).',
        'Thymeleaf Documentation [Электронный ресурс]. – Режим доступа: https://www.thymeleaf.org/documentation.html (дата обращения: 15.12.2024).',
        'MySQL 8.0 Reference Manual [Электронный ресурс]. – Режим доступа: https://dev.mysql.com/doc/refman/8.0/en/ (дата обращения: 15.12.2024).',
        'Hibernate ORM Documentation [Электронный ресурс]. – Режим доступа: https://hibernate.org/orm/documentation/ (дата обращения: 15.12.2024).',
        'Bloch, J. Effective Java / J. Bloch. – 3rd Edition. – Addison-Wesley, 2018. – 416 p.',
        'Horstmann, C.S. Core Java Volume I – Fundamentals / C.S. Horstmann. – 12th Edition. – Pearson, 2021. – 928 p.',
        'Schildt, H. Java: The Complete Reference / H. Schildt. – 12th Edition. – McGraw-Hill Education, 2021. – 1280 p.',
        'Методические материалы по выполнению курсовой работы по дисциплине «Современные технологии программирования» [Электронный ресурс]. – Режим доступа: https://palchevsky.ru/uploads/mm_stp_21.pdf (дата обращения: 15.12.2024).',
    ]
    
    for i, ref in enumerate(references, 1):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = p.add_run(f'{i}. {ref}')
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
        p.paragraph_format.first_line_indent = Cm(1.25)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    
    doc.add_page_break()


def add_appendix(doc):
    """Приложения"""
    add_heading1(doc, 'ПРИЛОЖЕНИЯ')
    
    # Приложение А
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Приложение А – Фрагмент исходного класса SecurityConfig')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run.bold = True
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)
    
    code_security = '''@Configuration
public class SecurityConfig {

    @Bean
    public PasswordEncoder passwordEncoder() {
        return new BCryptPasswordEncoder();
    }

    @Bean
    public UserDetailsService userDetailsService(UserRepository userRepository) {
        return username -> userRepository.findByUsername(username)
                .map(user -> org.springframework.security.core.userdetails.User
                        .withUsername(user.getUsername())
                        .password(user.getPassword())
                        .roles(user.getRole().name())
                        .build()
                )
                .orElseThrow(() -> new RuntimeException("User not found: " + username));
    }

    @Bean
    public SecurityFilterChain securityFilterChain(
            HttpSecurity http,
            DaoAuthenticationProvider authProvider
    ) throws Exception {
        http.authenticationProvider(authProvider);
        http
                .csrf(csrf -> csrf.disable())
                .authorizeHttpRequests(auth -> auth
                        .requestMatchers("/", "/about", "/schedule", "/tickets", 
                                "/radar", "/aircraft/**", "/login", "/register",
                                "/css/**", "/js/**", "/images/**", "/uploads/**"
                        ).permitAll()
                        .requestMatchers("/admin/home/**", "/admin/about/**")
                            .hasAnyRole("ADMIN", "MANAGER")
                        .requestMatchers("/booking/**").hasRole("USER")
                        .requestMatchers("/profile/**").hasRole("USER")
                        .requestMatchers("/admin/**").hasRole("ADMIN")
                        .anyRequest().permitAll()
                )
                .formLogin(form -> form
                        .loginPage("/login")
                        .defaultSuccessUrl("/", true)
                        .permitAll()
                )
                .logout(logout -> logout
                        .logoutSuccessUrl("/")
                        .permitAll()
                );
        return http.build();
    }
}'''
    
    for line in code_security.split('\n'):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(line)
        run.font.name = 'Courier New'
        run.font.size = Pt(10)
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        p.paragraph_format.space_after = Pt(0)
    
    doc.add_page_break()
    
    # Приложение Б
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Приложение Б – Фрагмент исходного класса BookingController')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run.bold = True
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)
    
    code_booking = '''@Controller
public class BookingController {

    private final FlightRepository flightRepository;
    private final UserRepository userRepository;
    private final BookingRepository bookingRepository;

    @GetMapping("/booking/confirm")
    public String confirmPage(@RequestParam("flightId") Long flightId,
                              @AuthenticationPrincipal UserDetails principal,
                              HttpServletRequest request,
                              Model model) {
        if (principal == null) return "redirect:/login";
        if (request.isUserInRole("ADMIN")) return "redirect:/";

        Flight flight = flightRepository.findById(flightId)
                .orElseThrow(() -> new RuntimeException("Flight not found"));
        User user = userRepository.findByUsername(principal.getUsername())
                .orElseThrow(() -> new RuntimeException("User not found"));

        if (bookingRepository.existsByUserAndFlight(user, flight)) {
            return "redirect:/profile?alreadyBooked";
        }
        if (flight.getAvailableSeats() == null || flight.getAvailableSeats() <= 0) {
            return "redirect:/tickets?noSeatsAvailable";
        }

        model.addAttribute("title", "Confirm booking");
        model.addAttribute("flight", flight);
        model.addAttribute("user", user);
        return "booking-confirm";
    }

    @PostMapping("/booking/confirm")
    @Transactional
    public String confirmBooking(@RequestParam("flightId") Long flightId,
                                 @RequestParam("passportNumber") String passportNumber,
                                 @RequestParam("address") String address,
                                 @AuthenticationPrincipal UserDetails principal,
                                 HttpServletRequest request) {
        // ... проверки ...
        
        Booking booking = new Booking(user, flight, LocalDateTime.now(), 
                                      passportNumber, address);
        bookingRepository.save(booking);
        flight.setAvailableSeats(flight.getAvailableSeats() - 1);
        flightRepository.save(flight);

        return "redirect:/profile?bookingSuccess";
    }
}'''
    
    for line in code_booking.split('\n'):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(line)
        run.font.name = 'Courier New'
        run.font.size = Pt(10)
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        p.paragraph_format.space_after = Pt(0)
    
    doc.add_page_break()
    
    # Приложение В
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Приложение В – Фрагмент исходного класса модели Flight')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run.bold = True
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)
    
    code_flight = '''@Entity
@Table(name = "flights")
public class Flight {

    @Id
    @GeneratedValue(strategy = GenerationType.IDENTITY)
    private Long id;

    @Column(name = "flight_number", nullable = false, length = 20)
    private String flightNumber;

    @Column(nullable = false, length = 100)
    private String airline;

    @Column(nullable = false, length = 100)
    private String origin;

    @Column(nullable = false, length = 100)
    private String destination;

    @DateTimeFormat(pattern = "yyyy-MM-dd'T'HH:mm")
    @Column(name = "departure_time", nullable = false)
    private LocalDateTime departureTime;

    @DateTimeFormat(pattern = "yyyy-MM-dd'T'HH:mm")
    @Column(name = "arrival_time", nullable = false)
    private LocalDateTime arrivalTime;

    @Column(name = "gate", length = 20)
    private String gate;

    @Column(name = "aircraft_type", length = 50)
    private String aircraftType;

    @Column(nullable = false, length = 20)
    private String status;

    @Column(name = "available_seats", nullable = false)
    private Integer availableSeats = 150;

    // Конструкторы, геттеры и сеттеры...
}'''
    
    for line in code_flight.split('\n'):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(line)
        run.font.name = 'Courier New'
        run.font.size = Pt(10)
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        p.paragraph_format.space_after = Pt(0)


def main():
    """Главная функция создания документа"""
    print("Создание документа...")
    
    doc = create_document()
    
    # Титульная страница
    add_title_page(doc)
    
    # Оглавление
    add_toc(doc)
    
    # Введение
    add_introduction(doc)
    
    # Раздел 1. Описание программы
    add_section1(doc)
    add_section1_2(doc)
    add_section1_3(doc)
    add_section1_interface(doc)
    add_section1_interface_continued(doc)
    
    # Раздел 2. Архитектура приложения
    add_section2(doc)
    
    # Раздел 3. Структура классов
    add_section3(doc)
    
    # Заключение
    add_conclusion(doc)
    
    # Список использованных источников
    add_references(doc)
    
    # Приложения
    add_appendix(doc)
    
    # Сохранение документа
    doc.save('Пояснительная_записка.docx')
    print("Документ успешно создан: Пояснительная_записка.docx")
    
    # Вывод информации о плейсхолдерах для картинок
    print("\n" + "="*60)
    print("МЕСТА ДЛЯ ВСТАВКИ СКРИНШОТОВ:")
    print("="*60)
    placeholders = [
        (1, "Архитектура безопасности", "Схема Spring Security (можно нарисовать в draw.io или взять из интернета)"),
        (2, "Взаимодействие в паттерне MVC", "Диаграмма MVC (Model-View-Controller)"),
        (3, "Навигация сайта", "Скриншот шапки сайта с меню"),
        (4, "Футер", "Скриншот нижней части страницы"),
        (5, "Регистрация", "Скриншот формы регистрации"),
        (6, "Авторизация", "Скриншот формы входа"),
        (7, "Главная страница", "Скриншот главной страницы"),
        (8, "Расписание рейсов", "Скриншот таблицы рейсов"),
        (9, "Страница бронирования билетов", "Скриншот страницы поиска билетов"),
        (10, "Подтверждение бронирования", "Скриншот формы подтверждения"),
        (11, "Профиль пользователя", "Скриншот профиля"),
        (12, "Редактирование профиля", "Скриншот формы редактирования профиля"),
        (13, "Каталог воздушных судов", "Скриншот галереи самолётов"),
        (14, "Детальная информация о воздушном судне", "Скриншот страницы самолёта"),
        (15, "Административная панель", "Скриншот админ-функций на странице расписания"),
        (16, "Форма редактирования рейса", "Скриншот формы добавления/редактирования рейса"),
        (17, "Форма редактирования воздушного судна", "Скриншот формы добавления/редактирования самолёта"),
        (18, "Фрагмент кода зависимостей", "Скриншот pom.xml из IDE"),
        (19, "Файл настроек подключения", "Скриншот application.properties из IDE"),
        (20, "Структура классов", "Скриншот структуры пакетов в IDE (Project view)"),
        (21, "Структура клиентской части", "Скриншот папки templates и static в IDE"),
    ]
    
    for num, title, desc in placeholders:
        print(f"\nРисунок {num} – {title}")
        print(f"   → {desc}")
    
    print("\n" + "="*60)
    print("НЕ ЗАБУДЬ ЗАМЕНИТЬ:")
    print("="*60)
    print("• [ГРУППА] — на номер твоей группы")
    print("• [ФИО СТУДЕНТА] — на твоё ФИО")
    print("="*60)

if __name__ == "__main__":
    main()
